import fitz
from PIL import Image
import pytesseract
import io
import re
import magic
from docx import Document
import subprocess
import csv


def extract_text_from_image(image_path):
    """
    Extract text from an image file using pytesseract.
    """
    try:
        image = Image.open(image_path)
        text = pytesseract.image_to_string(image)
        return text
    except Exception as e:
        raise RuntimeError(f"Failed to extract text from image: {str(e)}")

def extract_text_from_pdf(pdf_path):
    """
    Extract text from each page of a PDF document.
    """
    try:
        text = []
        pdf_document = fitz.open(pdf_path)
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            text.append(page.get_text())
            images = page.get_images(full=True)
            for img_index, img in enumerate(images):
                xref = img[0]
                base_image = pdf_document.extract_image(xref)
                image_bytes = base_image["image"]
                image = Image.open(io.BytesIO(image_bytes))
                text.append(pytesseract.image_to_string(image))
        pdf_document.close()
        return "\n".join(text)
    except Exception as e:
        raise RuntimeError(f"Failed to extract text from PDF: {str(e)}")

def extract_text_from_docx(docx_path):
    """
    Extract text from a DOCX document.
    """
    try:
        doc = Document(docx_path)
        text = [para.text for para in doc.paragraphs]
        return "\n".join(text)
    except Exception as e:
        raise RuntimeError(f"Failed to extract text from DOCX file: {str(e)}")

def extract_text_from_doc(doc_path):
    """
    Extract text from a DOC file using antiword for conversion.
    """
    try:
        result = subprocess.run(['antiword', doc_path], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        if result.returncode != 0:
            raise RuntimeError(result.stderr)
        return result.stdout
    except Exception as e:
        raise RuntimeError(f"Failed to extract text from DOC file: {str(e)}")

def extract_text_from_txt(txt_path):
    """
    Extract text from a TXT file.
    """
    try:
        with open(txt_path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        raise RuntimeError(f"Failed to extract text from TXT file: {str(e)}")
    

def extract_text_from_csv(file_path):
    """
    Extract text from a CSV file.
    """
    try:
        with open(file_path, newline='', encoding='utf-8') as csvfile:
            reader = csv.reader(csvfile)
            return '\n'.join([','.join(row) for row in reader])
    except Exception as e:
        raise RuntimeError(f"Error reading CSV file: {str(e)}")

def extract_emails(text):
    """
    Extract emails from text using regex.
    """
    try:
        email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
        return email_pattern.findall(text)
    except Exception as e:
        raise RuntimeError(f"Failed to extract emails from text: {str(e)}")

def extract_text_from_file(file_path):
    """
    Determine file type and extract text accordingly.
    """
    result = {
        'text': None,
        'error': None
    }

    try:
        mime = magic.Magic(mime=True)
        file_type = mime.from_file(file_path)

        if 'pdf' in file_type:
            result['text'] = extract_text_from_pdf(file_path)
        elif 'officedocument.wordprocessingml.document' in file_type:
            result['text'] = extract_text_from_docx(file_path)
        elif 'msword' in file_type or file_path.endswith('.doc'):
            result['text'] = extract_text_from_doc(file_path)
        elif 'image' in file_type:
            result['text'] = extract_text_from_image(file_path)
        elif 'text/plain' in file_type:
            result['text'] = extract_text_from_txt(file_path)
        elif 'csv' in file_type or file_path.endswith('.csv'):
            result['text'] = extract_text_from_csv(file_path)
        else:
            result['error'] = f"Unsupported file type: {file_type}"
    except Exception as e:
        result['error'] = f"Error extracting text from file: {str(e)}"
        result['text'] = ""  # Ensure 'text' is a string even if there is an error

    return result

def ocr_document(file_path):
    """
    Perform OCR on the given document file.
    """
    try:
        result = extract_text_from_file(file_path)
        emails = extract_emails(result['text'])
        return result, emails
    except Exception as e:
        raise RuntimeError(f"OCR failed for document: {str(e)}")

if __name__ == "__main__":
    document_path = 'DEVLOGIC Experience letter.docx'
    
    try:
        text, emails = ocr_document(document_path)
        print("Extracted Text:\n", text)
        print("Extracted Emails:\n", emails)
    except Exception as e:
        print(f"Error processing document: {str(e)}")


# import fitz 
# from PIL import Image
# import pytesseract
# import io
# import re
# import magic
# from docx import Document
# import subprocess
# import string
# import nltk
# from nltk.corpus import stopwords
# from nltk.tokenize import word_tokenize


# nltk.download('punkt')
# nltk.download('stopwords')

# def extract_text_from_image(image_path):
#     """
#     Extract text from an image file using pytesseract.
#     """
#     try:
#         image = Image.open(image_path)
#         text = pytesseract.image_to_string(image)
#         return text
#     except Exception as e:
#         raise RuntimeError(f"Failed to extract text from image: {str(e)}")

# def extract_text_from_pdf(pdf_path):
#     """
#     Extract text from each page of a PDF document.
#     """
#     try:
#         text = []
#         pdf_document = fitz.open(pdf_path)
#         for page_num in range(len(pdf_document)):
#             page = pdf_document.load_page(page_num)
#             text.append(page.get_text())
#             images = page.get_images(full=True)
#             for img_index, img in enumerate(images):
#                 xref = img[0]
#                 base_image = pdf_document.extract_image(xref)
#                 image_bytes = base_image["image"]
#                 image = Image.open(io.BytesIO(image_bytes))
#                 text.append(pytesseract.image_to_string(image))
#         pdf_document.close()
#         return "\n".join(text)
#     except Exception as e:
#         raise RuntimeError(f"Failed to extract text from PDF: {str(e)}")

# def extract_text_from_docx(docx_path):
#     """
#     Extract text from a DOCX document.
#     """
#     try:
#         doc = Document(docx_path)
#         text = [para.text for para in doc.paragraphs]
#         return "\n".join(text)
#     except Exception as e:
#         raise RuntimeError(f"Failed to extract text from DOCX file: {str(e)}")

# def extract_text_from_doc(doc_path):
#     """
#     Extract text from a DOC file using antiword for conversion.
#     """
#     try:
#         result = subprocess.run(['antiword', doc_path], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
#         if result.returncode != 0:
#             raise RuntimeError(result.stderr)
#         return result.stdout
#     except Exception as e:
#         raise RuntimeError(f"Failed to extract text from DOC file: {str(e)}")

# def extract_text_from_txt(txt_path):
#     """
#     Extract text from a TXT file.
#     """
#     try:
#         with open(txt_path, 'r', encoding='utf-8') as file:
#             return file.read()
#     except Exception as e:
#         raise RuntimeError(f"Failed to extract text from TXT file: {str(e)}")

# def clean_text(text):
#     """
#     Clean extracted text by removing special characters, stop words, and extra whitespace.
#     """
#     # Convert to lowercase
#     text = text.lower()

#     # Remove special characters and digits
#     text = re.sub(r'[^a-z\s]', '', text)

#     # Tokenize and remove stop words
#     words = word_tokenize(text)
#     stop_words = set(stopwords.words('english'))
#     filtered_words = [word for word in words if word not in stop_words]

#     # Join words back into a single string
#     cleaned_text = ' '.join(filtered_words)

#     return cleaned_text

# def extract_emails(text):
#     """
#     Extract emails from text using regex.
#     """
#     try:
#         email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
#         return email_pattern.findall(text)
#     except Exception as e:
#         raise RuntimeError(f"Failed to extract emails from text: {str(e)}")

# def extract_text_from_file(file_path):
#     """
#     Determine file type and extract text accordingly.
#     """
#     try:
#         mime = magic.Magic(mime=True)
#         file_type = mime.from_file(file_path)

#         if 'pdf' in file_type:
#             return extract_text_from_pdf(file_path)
#         elif 'officedocument.wordprocessingml.document' in file_type:
#             return extract_text_from_docx(file_path)
#         elif 'msword' in file_type or file_path.endswith('.doc'):
#             return extract_text_from_doc(file_path)
#         elif 'image' in file_type:
#             return extract_text_from_image(file_path)
#         elif 'text/plain' in file_type:
#             return extract_text_from_txt(file_path)
#         else:
#             raise ValueError(f"Unsupported file type: {file_type}")
#     except Exception as e:
#         raise RuntimeError(f"Error extracting text from file: {str(e)}")

# def ocr_document(file_path):
#     """
#     Perform OCR on the given document file and clean the extracted text.
#     """
#     try:
#         text = extract_text_from_file(file_path)
#         cleaned_text = clean_text(text)
#         emails = extract_emails(cleaned_text)
#         return cleaned_text, emails
#     except Exception as e:
#         raise RuntimeError(f"OCR failed for document: {str(e)}")

# if __name__ == "__main__":
#     document_path = 'Transcript.pdf'  
#     try:
#         text, emails = ocr_document(document_path)
#         print("Extracted and Cleaned Text:\n", text)
#         print("Extracted Emails:\n", emails)
#     except Exception as e:
#         print(f"Error processing document: {str(e)}")

